# app/docx_utils.py
from docx import Document
from pathlib import Path
import re

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    paragraphs = []
    for p in doc.paragraphs:
        paragraphs.append(p.text)
    return "\n".join(paragraphs)

def save_marked_docx(original_path: str, marks: list, out_path: str):
    """
    marks: list of dicts: {"match_text": "...", "comment": "...", "severity": "..."}
    We'll add inline reviews after paragraphs where match_text occurs.
    (python-docx doesn't support Word comment nodes well; this uses inline run insertion.)
    """
    doc = Document(original_path)
    for p in doc.paragraphs:
        for m in marks:
            if m["match_text"].strip() and m["match_text"].strip() in p.text:
                # Append a visible inline comment marker
                marker = f"\n\n<<REVIEW - {m.get('severity','Info')}>> {m.get('comment')}"
                p.add_run(marker)
    p_out = Path(out_path)
    p_out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path

def find_paragraph_matches(text: str, patterns):
    """
    patterns: list of regex patterns or strings
    Returns list of matches with snippet and approximate location text
    """
    res = []
    for pat in patterns:
        if pat.startswith("re:"):
            regex = pat[3:]
            for m in re.finditer(regex, text, flags=re.I):
                res.append({"match_text": m.group(0), "span": m.span()})
        else:
            idx = text.lower().find(pat.lower())
            if idx != -1:
                res.append({"match_text": pat, "span": (idx, idx+len(pat))})
    return res
